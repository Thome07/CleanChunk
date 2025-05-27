# bom, porem tem Respostas que estao vindo muito grandes


from tkinter import *
from tkinter import ttk, filedialog, messagebox
from sentence_transformers import SentenceTransformer
import numpy as np
import re
import os
import json
import pdfplumber
import docx
from unidecode import unidecode
import ftfy
import logging
from typing import List, Dict, Any, Optional
import threading
from queue import Queue
import google.generativeai as genai
import time
import requests
from datetime import datetime, timedelta

try:
    from transformers import pipeline, AutoTokenizer, AutoModel
    import torch
    from sklearn.cluster import KMeans
    from sklearn.metrics.pairwise import cosine_similarity
    import umap
    semantic_analyzer = pipeline("text-classification", model="microsoft/DialoGPT-medium")
    coherence_checker = pipeline("text-classification", model="textattack/bert-base-uncased-CoLA")
    ADVANCED_NLP = True
except ImportError:
    ADVANCED_NLP = False
    print("AVISO: Bibliotecas avançadas não instaladas. Usando fallback.")

class DocumentProcessor:
    def __init__(self, gemini_api_key: str = None):
        # Modelo mais leve e rápido, mas ainda eficiente
        self.model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
        self.model.max_seq_length = 256
        
        # Configuração Gemini
        self.gemini_api_key = gemini_api_key
        self.gemini_model = None
        self.use_gemini = False
        self.request_count = 0
        self.last_request_time = datetime.now()
        
        # Parâmetros otimizados para chunks SEMÂNTICOS (não fixos!)
        self.min_chunk_words = 40      # Mínimo em palavras
        self.max_chunk_words = 150     # Máximo em palavras  
        self.ideal_chunk_words = 80    # Ideal em palavras
        self.min_sentences_per_chunk = 2
        self.max_sentences_per_chunk = 8
        
        # Rate limiting otimizado
        self.max_requests_per_minute = 14  # Margem de segurança
        self.request_interval = 4.0        # 4.0 segundos entre requests
        self.batch_size_gemini = 15         # Processa 15 sentenças por vez com Gemini
        
        if gemini_api_key:
            try:
                genai.configure(api_key=gemini_api_key)
                self.gemini_model = genai.GenerativeModel('gemini-1.5-flash')
                self.use_gemini = True
                print("✅ Gemini AI configurado com sucesso!")
            except Exception as e:
                print(f"❌ Erro ao configurar Gemini: {e}")
                self.use_gemini = False
        
        # Threshold semântico mais flexível
        self.threshold_base = 0.75
        self.threshold_flexible = 0.65  # Para chunks pequenos
        self.batch_size = 32
        
        # Cache para otimização
        self._validation_cache = {}
        self._embedding_cache = {}
        self._semantic_cache = {}
        
        # Modelo maior para melhor qualidade semântica
        self.model_semantica = SentenceTransformer('all-MiniLM-L12-v2')  # Modelo mais preciso
        self.threshold_semantica = 0.85  # Threshold mais alto para maior precisão
        
        # Parâmetros otimizados para semântica perfeita
        self.min_words = 15  # Mínimo em palavras, não caracteres
        self.max_words = 80  # Máximo em palavras (60-120 como sugerido)
        self.ideal_words = 50  # Tamanho ideal
        self.min_sentences = 2  # Mínimo de frases por chunk
        self.max_sentences = 6  # Máximo de frases por chunk
        
        # Regex melhorados
        self.sentence_splitter = re.compile(r'(?<=[.!?])\s+(?=[A-Z])')
        self.paragraph_splitter = re.compile(r'\n\s*\n')
        
        # Padrões para detectar ideias completas
        self.idea_starters = re.compile(r'\b(?:Portanto|Assim|Dessa forma|Por isso|Consequentemente|Em suma|Além disso|Também|Primeiro|Segundo|Finalmente)\b', re.IGNORECASE)
        self.idea_enders = re.compile(r'\b(?:concluindo|resumindo|em conclusão|por fim|finalmente)\b', re.IGNORECASE)

        # Regex patterns necessários
        self.page_re = re.compile(r'\[Página\s+(\d+)\]')
        self.header_re = re.compile(r'^[A-Z\s]{5,}$')
        self.faq_q_re = re.compile(r'^\d+\.\s*(.+)$')
        self.faq_a_re = re.compile(r'^R:\s*(.+)$')
        self.header_interview_re = re.compile(r'^[A-Z\s]{5,}$')
        self.speaker_re = re.compile(r'^([^:]+):\s*(.+)$')

        self._gemini_cache = {}

    def tem_sentido(self, texto: str) -> bool:
        """Verifica se o texto tem sentido usando modelo de gramática"""
        if not coherence_checker or len(texto) < 3:
            return len(texto) >= 10  # Fallback mais permissivo
        try:
            resultado = coherence_checker(texto[:512])
            return resultado[0]['label'] == 'LABEL_1' and resultado[0]['score'] > 0.6  # Threshold mais baixo
        except Exception:
            return len(texto) >= 10

    def limpar_texto(self, texto: str) -> str:
        """Limpeza otimizada de texto"""
        if not texto:
            return ""
        
        # Cache para textos já limpos
        if texto in self._validation_cache:
            return self._validation_cache[texto]
        
        # Fix encoding issues
        texto = ftfy.fix_text(texto)
        
        # Limpeza em uma passada usando regex
        replacements = [
            (r'\\+["\']', '"'),
            (r'-\s*\n', ''),
            (r'\r\n|\r|\n', ' '),
            (r'\s+', ' '),
            (r'\s+([,\.!?;:])', r'\1'),
            (r'\\(["\'\\])', r'\1'),
            (r'\\n', ' '),
            (r'\u201c|\u201d|&quot;', '"'),
            (r'\u2018|\u2019|&#39;', "'"),
        ]
        
        for pattern, replacement in replacements:
            texto = re.sub(pattern, replacement, texto)
        
        # Cache do resultado
        result = texto.strip()
        self._validation_cache[texto] = result
        return result

    def preprocess_text(self, text: str) -> str:
        """Preprocessamento específico para limpeza de ruído"""
        if not text:
            return ""
            
        # Remove linhas numéricas isoladas
        text = re.sub(r'^\s*\d+\.?\s*$', '', text, flags=re.MULTILINE)
        # Remove palavras em CAPS muito curtas
        text = re.sub(r'\b[A-Z]{1,4}\b\.?', '', text)
        # Remove pontuação repetida
        text = re.sub(r'([.!?])\1+', r'\1', text)
        
        return text.strip()

    def load_document(self, path: str) -> str:
        """Carrega documento de diferentes formatos"""
        if not os.path.exists(path):
            raise FileNotFoundError(f"Arquivo não encontrado: {path}")
            
        ext = os.path.splitext(path)[1].lower()
        content = ""
        
        try:
            if ext == '.pdf':
                with pdfplumber.open(path) as pdf:
                    parts = []
                    for i, page in enumerate(pdf.pages, 1):
                        raw = page.extract_text()
                        if raw:
                            raw = raw.replace('\\"', '"').replace("\\'", "'")
                        if raw.strip():  # Só adiciona se tem conteúdo
                            parts.append(f"[Página {i}]\n{raw}")
                    content = "\n\n".join(parts)
            elif ext == '.docx':
                doc = docx.Document(path)
                parts = [para.text for para in doc.paragraphs if para.text.strip()]
                content = "\n\n".join(parts)
            else:  # .txt
                with open(path, 'r', encoding='utf-8') as f:
                    content = f.read()
        except Exception as e:
            raise Exception(f"Erro ao ler arquivo: {str(e)}")
        
        return content

    def load_lines(self, path: str, is_faq_mode: bool = False) -> List[str]: # Adicionado is_faq_mode
        """Carrega documento preservando linhas, com tratamento especial para modo FAQ."""
        if not os.path.exists(path):
            raise FileNotFoundError(f"Arquivo não encontrado: {path}")
            
        ext = os.path.splitext(path)[1].lower()
        lines = []
        
        try:
            if ext == '.pdf':
                with pdfplumber.open(path) as pdf:
                    for i, page in enumerate(pdf.pages, 1):
                        raw = page.extract_text(layout=True) 
                        if raw:
                            raw = raw.replace('-\n', '')  # Remove hifens de quebra de linha

                            if is_faq_mode:
                                # Para FAQ, mantenha as linhas individuais da página
                                page_actual_lines = [l.strip() for l in raw.splitlines() if l.strip()]
                                lines.extend(page_actual_lines)
                            else:
                                # Comportamento original para outros modos (não FAQ)
                                raw = raw.replace('\n', ' ') # Substitui quebras por espaços
                                if raw.strip(): # Só adiciona se tem conteúdo
                                    lines.append(f"[Página {i}]")
                                    # Após replace('\n', ' '), raw.splitlines() provavelmente terá 1 item ou poucos
                                    lines.extend([l.strip() for l in raw.splitlines() if l.strip()])
            elif ext == '.docx':
                doc = docx.Document(path)
                if is_faq_mode:
                    for para in doc.paragraphs:
                        if para.text.strip(): # Checa se o parágrafo tem algum conteúdo
                            # Para FAQ, divide o parágrafo por quebras de linha internas
                            paragraph_lines = [line.strip() for line in para.text.splitlines() if line.strip()]
                            lines.extend(paragraph_lines)
                else:
                    # Comportamento original para não-FAQ: cada parágrafo não vazio é uma linha
                    for para in doc.paragraphs:
                        if para.text.strip():
                             lines.append(para.text) # O original mantinha o texto do parágrafo
            elif ext == '.txt': # .txt
                with open(path, 'r', encoding='utf-8') as f:
                    # Para TXT, o comportamento de preservar linhas já é adequado para FAQ
                    lines = [l.strip() for l in f if l.strip()] # strip() em vez de rstrip('\n') para remover espaços em branco das extremidades
            else:
                 # Caso para outros tipos de arquivo ou se um comportamento padrão for necessário
                 # Você pode querer levantar um erro para tipos não suportados ou ter um fallback
                with open(path, 'r', encoding='utf-8') as f:
                    lines = [l.strip() for l in f if l.strip()]

        except Exception as e:
            # Adiciona o modo FAQ à mensagem de erro se aplicável
            error_prefix = f"Erro ao ler arquivo para {'FAQ' if is_faq_mode else 'processamento normal'}: "
            raise Exception(error_prefix + str(e))
        
        return lines
    
    def split_sentences(self, text: str) -> List[str]:
        """Divide texto em sentenças"""
        if not text:
            return []
        # Regex melhorado para lidar com diferentes casos de pontuação
        sentences = re.split(r'(?<=[.!?])\s+(?=[A-ZÀ-Ü])', text)  # Adicionado suporte para caracteres acentuados
        return [s.strip() for s in sentences if s.strip()]

    def _is_chunk_valid(self, chunk_sentences: List[str]) -> bool:
        """Valida se um chunk tem conteúdo semântico válido"""
        if not chunk_sentences:
            return False
            
        chunk_text = ' '.join(chunk_sentences).strip()
        
        # Critérios básicos de tamanho
        if len(chunk_text) < self.min_words:
            return False
        
        # Verifica se tem palavras suficientes (não só pontuação/números)
        words = re.findall(r'\b\w+\b', chunk_text)
        if len(words) < 10:  # Pelo menos 10 palavras
            return False
        
        # Verifica se não termina no meio de uma frase
        if not chunk_text.rstrip().endswith(('.', '!', '?', ':', ';')):
            # Se não termina com pontuação, verifica se a última sentença está completa
            last_sentence = chunk_sentences[-1].strip()
            if len(last_sentence) < 20 or not self._sentence_seems_complete(last_sentence):
                return False
        
        # Verifica se não tem muitos caracteres estranhos ou fragmentação
        clean_chars = re.sub(r'[^\w\s]', '', chunk_text)
        if len(clean_chars) < len(chunk_text) * 0.7:  # Muito símbolo/pontuação
            return False
        
        # Verifica diversidade vocabular básica
        unique_words = set(word.lower() for word in words if len(word) > 2)
        if len(unique_words) < len(words) * 0.4:  # Muito repetitivo
            return False
        
        # Verifica se não é só números ou códigos
        if re.match(r'^[\d\s\-\.]+$', chunk_text):
            return False
            
        return True

    def _sentence_seems_complete(self, sentence: str) -> bool:
        """Verifica se uma sentença parece completa (método já estava sendo chamado)"""
        if not sentence:
            return False
        
        sentence = sentence.strip()
        
        # Muito curta
        if len(sentence) < 15:
            return False
        
        # Termina com pontuação final
        if sentence.endswith(('.', '!', '?')):
            return True
        
        # Termina com dois pontos ou ponto-e-vírgula (pode ser válido)
        if sentence.endswith((':', ';')):
            return len(sentence) > 25  # Só aceita se for longa o suficiente
        
        # Verifica se não termina com conectivos ou preposições
        incomplete_endings = [
            r'\b(e|mas|que|quando|se|porque|para|com|em|de|da|do|na|no|pela|pelo|sobre|entre|durante|antes|depois)\s*$',
            r'\b(o|a|os|as|um|uma|uns|umas)\s*$',  # Artigos soltos
            r'\b(muito|mais|menos|bem|mal|melhor|pior)\s*$',  # Advérbios incompletos
        ]
        
        for pattern in incomplete_endings:
            if re.search(pattern, sentence, re.IGNORECASE):
                return False
        
        # Se chegou até aqui e tem tamanho razoável, provavelmente está completa
        return len(sentence) >= 20

    def _clean_chunk_text(self, text: str) -> str:
        """Versão otimizada da limpeza de texto para chunks"""
        if not text:
            return ""
        
        # Cache para evitar processamento repetido
        cache_key = hash(text[:100])  # Use primeiros 100 chars como chave
        if cache_key in self._validation_cache:
            cached_result = self._validation_cache.get(f"clean_{cache_key}")
            if cached_result is not None:
                return cached_result
        
        # Limpeza otimizada em uma passada
        # Remove quebras de linha e normaliza espaços
        text = re.sub(r'\s+', ' ', text)
        
        # Remove caracteres de escape problemáticos
        text = text.replace('\\n', ' ').replace('\\t', ' ').replace('\\r', ' ')
        text = re.sub(r'\\+', '', text)
        
        # Corrige pontuação
        text = re.sub(r'([.!?])\s*\1+', r'\1', text)  # Remove pontuação duplicada
        text = re.sub(r'\s+([,.!?;:])', r'\1', text)  # Remove espaços antes de pontuação
        text = re.sub(r'([,.!?;:])\s*([,.!?;:])', r'\1 \2', text)  # Espaça pontuação múltipla
        
        # Limpa início e fim
        text = re.sub(r'^[^\w\s]*', '', text)  # Remove símbolos no início
        text = re.sub(r'[^\w\s.!?]*$', '', text)  # Remove símbolos no fim (exceto pontuação)
        
        # Garante que termine com pontuação se não terminar
        text = text.strip()
        if text and not text[-1] in '.!?:;':
            # Só adiciona ponto se a última palavra parecer completa
            words = text.split()
            if words and len(words[-1]) >= 3:
                text += '.'
        
        # Cache do resultado
        self._validation_cache[f"clean_{cache_key}"] = text
        
        return text

    def _rate_limit_gemini(self):
        """Rate limiting otimizado e não-bloqueante"""
        now = datetime.now()
        time_diff = (now - self.last_request_time).total_seconds()
        
        # Reset contador a cada minuto
        if time_diff >= 60:
            self.request_count = 0
            self.last_request_time = now
        
        # Se atingiu limite, calcula tempo de espera preciso
        if self.request_count >= self.max_requests_per_minute:
            wait_time = 60 - time_diff
            if wait_time > 0:
                print(f"⏳ Rate limit: aguardando {wait_time:.1f}s...")
                time.sleep(wait_time + 0.5)  # +0.5s margem
                self.request_count = 0
                self.last_request_time = datetime.now()
        
        # Intervalo mínimo otimizado
        elif time_diff < self.request_interval:
            sleep_time = self.request_interval - time_diff
            time.sleep(sleep_time)

    def _gemini_analyze_semantic_break_batch(self, sentences: List[str]) -> List[bool]:
        cache_key = hash(tuple(sentences))
        if cache_key in self._gemini_cache:
            return self._gemini_cache[cache_key]
        if not self.use_gemini or len(sentences) < 2:
            return [False] * (len(sentences) - 1)
        try:
            self._rate_limit_gemini()
            # Prepara texto das sentenças
            sentence_pairs = []
            for i in range(len(sentences) - 1):
                s1 = sentences[i][:150]  # Limita tamanho
                s2 = sentences[i + 1][:150]
                sentence_pairs.append(f"{i+1}. \"{s1}\" → \"{s2}\"")
            max_pairs = 20  # Aumente o número de pares analisados por request
            if len(sentence_pairs) > max_pairs:
                result = self._gemini_analyze_semantic_break_batch(sentences[:max_pairs+1]) + \
                         self._gemini_analyze_semantic_break_batch(sentences[max_pairs:])
                self._gemini_cache[cache_key] = result
                return result
            pairs_text = "\n".join(sentence_pairs)
            prompt = f"""Analise as transições entre estas {len(sentence_pairs)} sentenças consecutivas.\nPara cada par, responda 1 (QUEBRA) ou 0 (CONTINUA) separados por vírgula:\n\n{pairs_text}\n\nResposta:"""
            response = self.gemini_model.generate_content(prompt)
            self.request_count += 1
            # Parse da resposta
            breaks_text = response.text.strip()
            breaks = []
            for x in breaks_text.split(','):
                try:
                    breaks.append(int(x.strip()) == 1)
                except:
                    breaks.append(False)
            # Ajusta tamanho se necessário
            expected_len = len(sentences) - 1
            if len(breaks) != expected_len:
                breaks = breaks[:expected_len] + [False] * (expected_len - len(breaks))
            self._gemini_cache[cache_key] = breaks
            return breaks
        except Exception as e:
            print(f"⚠️ Erro Gemini batch: {e}")
            return [False] * (len(sentences) - 1)

    def chunk_semantic_pairwise(self, sentences: List[str], pagina: Optional[int] = None) -> List[Dict]:
        """Chunking semântico ADAPTATIVO com foco na qualidade semântica"""
        if not sentences:
            return []
        
        # Filtra e limpa sentenças
        valid_sentences = []
        for sentence in sentences:
            cleaned = self.limpar_texto(sentence)
            if len(cleaned) >= 15 and len(re.findall(r'\b\w+\b', cleaned)) >= 3:
                valid_sentences.append(cleaned)
        
        if len(valid_sentences) < 2:
            return []
        
        chunks = []
        print(f"📝 Processando {len(valid_sentences)} sentenças...")
        
        if self.use_gemini:
            print("🤖 Usando Gemini AI para análise semântica...")
            chunks = self._chunk_with_gemini_semantic(valid_sentences, pagina)
        else:
            print("📊 Usando embeddings para análise semântica...")
            chunks = self._chunk_with_embeddings_semantic(valid_sentences, pagina)
        
        print(f"✅ Gerados {len(chunks)} chunks semânticos")
        return chunks

    def _chunk_with_gemini_semantic(self, sentences: List[str], pagina: Optional[int] = None) -> List[Dict]:
        """Chunking com Gemini focado na semântica"""
        chunks = []
        current_chunk = []
        
        # Processa em lotes menores para eficiência
        batch_size = self.batch_size_gemini
        i = 0
        
        while i < len(sentences):
            # Pega lote atual
            end_idx = min(i + batch_size, len(sentences))
            batch_sentences = sentences[i:end_idx]
            
            # Analisa quebras semânticas do lote
            if len(batch_sentences) > 1:
                semantic_breaks = self._gemini_analyze_semantic_break_batch(batch_sentences)
            else:
                semantic_breaks = []
            
            # Processa sentenças do lote
            for j, sentence in enumerate(batch_sentences):
                current_chunk.append(sentence)
                word_count = self._count_words_in_chunk(current_chunk)
                
                should_break = False
                
                # Critérios para quebra
                if word_count >= self.max_chunk_words:
                    should_break = True
                elif (word_count >= self.ideal_chunk_words and 
                      len(current_chunk) >= self.min_sentences_per_chunk):
                    # Verifica se Gemini sugere quebra semântica
                    if j < len(semantic_breaks) and semantic_breaks[j]:
                        should_break = True
                elif len(current_chunk) >= self.max_sentences_per_chunk:
                    should_break = True
                
                # Cria chunk se deve quebrar
                if should_break and len(current_chunk) > 1:
                    # Remove última sentença para o próximo chunk
                    chunk_sentences = current_chunk[:-1]
                    if self._is_valid_semantic_chunk_words(chunk_sentences):
                        chunk_text = self._clean_chunk_text_optimized(' '.join(chunk_sentences))
                        if chunk_text:
                            chunks.append({
                                'pagina': pagina,
                                'conteudo': chunk_text,
                                'tipo_quebra': 'semantica_gemini',
                                'confianca': 0.92,
                                'palavras': len(chunk_text.split()),
                                'sentencas': len(chunk_sentences)
                            })
                    
                    # Inicia novo chunk com a sentença atual
                    current_chunk = [sentence]
            
            i = end_idx
            print(f"📊 Processado lote {i//batch_size}/{(len(sentences) + batch_size - 1)//batch_size}")
        
        # Processa último chunk
        if current_chunk and self._is_valid_semantic_chunk_words(current_chunk):
            chunk_text = self._clean_chunk_text_optimized(' '.join(current_chunk))
            if chunk_text:
                chunks.append({
                    'pagina': pagina,
                    'conteudo': chunk_text,
                    'tipo_quebra': 'final',
                    'confianca': 0.90,
                    'palavras': len(chunk_text.split()),
                    'sentencas': len(current_chunk)
                })
        
        return chunks

    def _chunk_with_embeddings_semantic(self, sentences: List[str], pagina: Optional[int] = None) -> List[Dict]:
        """Chunking com embeddings aprimorado com análise de coesão."""
        if not sentences:
            return []
            
        chunks = []
        current_chunk_sentences = []
        similarities = self._calculate_semantic_similarity_batch(sentences)
        
        i = 0
        while i < len(sentences):
            sentence = sentences[i]
            current_chunk_sentences.append(sentence)
            
            word_count = self._count_words_in_chunk(current_chunk_sentences)
            num_sentences = len(current_chunk_sentences)
            
            # Verifica se é a última sentença do parágrafo
            is_last_sentence = (i == len(sentences) - 1)
            
            # Define o limiar de similaridade para a quebra.
            # Um valor mais baixo significa que a próxima frase tem que ser bem diferente para causar uma quebra.
            similarity_threshold = 0.5 

            # Condição de quebra: a próxima frase tem baixa similaridade E o chunk atual já tem um tamanho razoável
            break_due_to_similarity = (i < len(similarities) and 
                                       similarities[i] < similarity_threshold and 
                                       word_count >= self.min_chunk_words)

            # Condição de quebra: o chunk está ficando muito grande
            break_due_to_max_size = word_count >= self.max_chunk_words

            # Se for a última sentença ou se uma condição de quebra foi atingida
            if is_last_sentence or break_due_to_similarity or break_due_to_max_size:
                
                # Validação final usando coesão
                cohesion_score = self._calculate_chunk_cohesion(current_chunk_sentences)
                
                # Só aceita o chunk se ele for coeso ou se for o final do parágrafo
                if cohesion_score > 0.7 or is_last_sentence:
                    chunk_text = self._clean_chunk_text_optimized(' '.join(current_chunk_sentences))
                    if self._is_valid_semantic_chunk_words(current_chunk_sentences):
                        chunks.append({
                            'pagina': pagina,
                            'conteudo': chunk_text,
                            'tipo_quebra': 'semantica_coesao' if not is_last_sentence else 'final_paragrafo',
                            'confianca': float(cohesion_score),
                            'palavras': len(chunk_text.split()),
                            'sentencas': len(current_chunk_sentences)
                        })
                    current_chunk_sentences = [] # Limpa para o próximo chunk
                # Se não for coeso, tenta adicionar a próxima frase para ver se melhora a coesão
            
            i += 1
            
        # Garante que qualquer resto no buffer seja processado
        if current_chunk_sentences:
            chunk_text = self._clean_chunk_text_optimized(' '.join(current_chunk_sentences))
            if self._is_valid_semantic_chunk_words(current_chunk_sentences):
                chunks.append({
                    'pagina': pagina,
                    'conteudo': chunk_text,
                    'tipo_quebra': 'final',
                    'confianca': 0.85,
                    'palavras': len(chunk_text.split()),
                    'sentencas': len(current_chunk_sentences)
                })

        return chunks

    def _count_words_in_chunk(self, sentences: List[str]) -> int:
        """Conta palavras no chunk atual"""
        text = ' '.join(sentences)
        return len(re.findall(r'\b\w+\b', text))

    def _is_valid_semantic_chunk_words(self, sentences: List[str]) -> bool:
        """Valida chunk baseado em contagem de palavras"""
        if not sentences:
            return False
        
        word_count = self._count_words_in_chunk(sentences)
        
        # Critérios baseados em palavras, não caracteres
        if word_count < self.min_chunk_words:
            return False
        
        text = ' '.join(sentences)
        
        # Verifica se termina adequadamente
        if not text.rstrip().endswith(('.', '!', '?', ':', ';')):
            if word_count < self.ideal_chunk_words:  # Só aceita se for chunk grande
                return False
        
        # Verifica diversidade lexical
        words = re.findall(r'\b\w{2,}\b', text.lower())
        unique_words = set(words)
        if len(words) > 0 and len(unique_words) / len(words) < 0.35:
            return False
        
        return True

    def chunk_structured(self, path: str) -> List[Dict]:
        """Versão otimizada que respeita parágrafos como unidades semânticas."""
        # Em vez de carregar por linhas, carregamos o documento inteiro para preservar parágrafos.
        full_text = self.load_document(path)
        all_chunks = []
        current_page = 1

        # Usa o regex de parágrafo para dividir o texto em unidades lógicas
        paragraphs = self.paragraph_splitter.split(full_text)
        print(f"📖 Documento com {len(paragraphs)} parágrafos detectados.")

        for paragraph_text in paragraphs:
            # Pula parágrafos vazios ou muito curtos
            if len(paragraph_text.strip()) < self.min_words:
                continue

            # Detecta se o parágrafo contém uma marcação de página
            page_match = self.page_re.search(paragraph_text)
            if page_match:
                current_page = int(page_match.group(1))
                # Remove a marcação de página do texto para não sujar os chunks
                paragraph_text = self.page_re.sub('', paragraph_text)

            # Detecta se o parágrafo é um cabeçalho
            if self.header_re.match(paragraph_text.strip()):
                # Trata o cabeçalho como um chunk separado e pequeno, se relevante
                # Ou pode ser usado como metadados para os próximos chunks
                # Aqui, vamos simplesmente pular para não virar um chunk de texto
                continue

            # Limpa e divide o parágrafo em sentenças
            cleaned_paragraph = self.limpar_texto(paragraph_text)
            sentences = self.split_sentences(cleaned_paragraph)

            if not sentences:
                continue
            
            # Agora, o chunking semântico opera DENTRO de um parágrafo,
            # o que o impede de juntar ideias de parágrafos diferentes.
            print(f"🔄 Processando parágrafo com {len(sentences)} sentenças (página {current_page})")
            
            # A função chunk_semantic_pairwise é chamada para cada parágrafo
            paragraph_chunks = self.chunk_semantic_pairwise(sentences, current_page)
            
            all_chunks.extend(paragraph_chunks)

        # O pós-processamento continua o mesmo
        final_chunks = self._post_process_semantic_chunks(all_chunks)
        
        print(f"🎯 Total: {len(final_chunks)} chunks semânticos")
        if final_chunks:
            words = [chunk.get('palavras', 0) for chunk in final_chunks]
            print(f"📏 Palavras - Mín: {min(words)}, Máx: {max(words)}, Média: {sum(words)/len(words):.0f}")
        
        return final_chunks

    def _calculate_chunk_cohesion(self, sentences: List[str]) -> float:
        """Calcula a coesão interna de um chunk medindo a similaridade média ao centroide."""
        if not sentences or not ADVANCED_NLP:
            return 0.0
        
        try:
            embeddings = self.model.encode(sentences, show_progress_bar=False)
            # Calcula o centroide (embedding médio) do chunk
            centroid = np.mean(embeddings, axis=0)
            
            # Calcula a similaridade de cada sentença com o centroide
            similarities = cosine_similarity(embeddings, [centroid])
            
            # A coesão é a similaridade média. Quanto maior, mais coeso é o chunk.
            return np.mean(similarities)
        except Exception as e:
            print(f"⚠️ Erro ao calcular coesão do chunk: {e}")
            return 0.0

    def _post_process_semantic_chunks(self, chunks: List[Dict]) -> List[Dict]:
        """Pós-processamento focado na qualidade semântica, incluindo divisão de chunks longos."""
        if not chunks:
            return []
        
        # 1. Limpeza inicial e filtragem básica de chunks recebidos
        #    Garante que 'conteudo', 'palavras', 'sentencas' estejam consistentes.
        pre_processed_candidates = []
        for chunk in chunks:
            content = chunk.get('conteudo', '')
            # Limpeza básica que não deve alterar drasticamente o conteúdo semântico
            content = re.sub(r'\s+', ' ', content).strip()
            content = re.sub(r'^[^\w]*(?=[a-zA-Z0-9À-Üà-ü])', '', content) # Remove lixo no início, preservando o primeiro caractere útil
            
            if not content:
                continue

            # Recalcula sentenças e palavras com base no conteúdo limpo
            current_sentences = self.split_sentences(content)
            if not current_sentences: # Pula se não houver sentenças após a limpeza
                continue
                
            word_count = self._count_words_in_chunk(current_sentences)

            # Atualiza o chunk com o conteúdo limpo e contagens recalculadas
            chunk['conteudo'] = content
            chunk['palavras'] = word_count
            chunk['sentencas'] = len(current_sentences)
            
            # Validação semântica e de tamanho mínimo
            if self._is_valid_semantic_chunk_words(current_sentences):
                pre_processed_candidates.append(chunk)
            # else:
                # print(f"ℹ️ Chunk descartado na pré-validação: {content[:60]}... ({word_count} palavras)")
        
        if not pre_processed_candidates:
            return []

        # 2. Otimiza fronteiras dos chunks (mesclando chunks pequenos adjacentes)
        merged_chunks = self._optimize_chunk_boundaries(pre_processed_candidates)
        if not merged_chunks: # Se a lista estiver vazia ou for None
            return [] 
        
        # Filtra explicitamente quaisquer Nones que possam ter escapado, embora não devessem.
        merged_chunks = [mc for mc in merged_chunks if mc is not None]
        if not merged_chunks: # Se a lista ficar vazia após a filtragem
            return []

        
        # 3. Divide chunks que ainda estão muito grandes após a mesclagem
        #    Esta função pode alterar o número de chunks e suas composições.
        split_oversized_chunks = self._split_all_oversized_chunks(merged_chunks)

        # 4. Processamento final: Atribuição de ID sequencial e cálculo de 'qualidade'
        final_output_chunks = []
        for i, chunk_data in enumerate(split_oversized_chunks, 1):
            # Revalida cada chunk final, pois as operações anteriores podem ter alterado sua validade
            final_content = chunk_data['conteudo']
            final_sentences = self.split_sentences(final_content) # Re-split para consistência

            if not final_sentences: # Segurança extra
                continue

            # Recalcula palavras e sentenças para máxima precisão antes do score de qualidade
            chunk_data['palavras'] = self._count_words_in_chunk(final_sentences)
            chunk_data['sentencas'] = len(final_sentences)

            # Só adiciona se o chunk final ainda for válido e tiver um tamanho mínimo
            if chunk_data['palavras'] >= self.min_chunk_words and \
               self._is_valid_semantic_chunk_words(final_sentences): # _is_valid_semantic_chunk_words já checa min_chunk_words
                
                chunk_data['chunk_id'] = i
                chunk_data['qualidade'] = self._calculate_chunk_quality_score(chunk_data) # Recalcula qualidade com dados finais
                final_output_chunks.append(chunk_data)
            # else:
            #     print(f"ℹ️ Chunk descartado no pós-processamento final: {final_content[:60]}... ({chunk_data['palavras']} palavras)")
            
        return final_output_chunks

    def _clean_chunk_text_optimized(self, text: str) -> str:
        """Versão otimizada da limpeza de texto para chunks"""
        if not text:
            return ""
        
        # Cache para evitar processamento repetido
        cache_key = hash(text[:100])  # Use primeiros 100 chars como chave
        if cache_key in self._validation_cache:
            cached_result = self._validation_cache.get(f"clean_{cache_key}")
            if cached_result is not None:
                return cached_result
        
        # Limpeza otimizada em uma passada
        # Remove quebras de linha e normaliza espaços
        text = re.sub(r'\s+', ' ', text)
        
        # Remove caracteres de escape problemáticos
        text = text.replace('\\n', ' ').replace('\\t', ' ').replace('\\r', ' ')
        text = re.sub(r'\\+', '', text)
        
        # Corrige pontuação
        text = re.sub(r'([.!?])\s*\1+', r'\1', text)  # Remove pontuação duplicada
        text = re.sub(r'\s+([,.!?;:])', r'\1', text)  # Remove espaços antes de pontuação
        text = re.sub(r'([,.!?;:])\s*([,.!?;:])', r'\1 \2', text)  # Espaça pontuação múltipla
        
        # Limpa início e fim
        text = re.sub(r'^[^\w\s]*', '', text)  # Remove símbolos no início
        text = re.sub(r'[^\w\s.!?]*$', '', text)  # Remove símbolos no fim (exceto pontuação)
        

        text = text.strip()
        if text and not text[-1] in '.!?:;':

            words = text.split()
            if words and len(words[-1]) >= 3:
                text += '.'
        
        # Cache do resultado
        self._validation_cache[f"clean_{cache_key}"] = text
        
        return text

    def _validate_chunk_content(self, content: str) -> bool:
        """Validação final do conteúdo do chunk"""
        if not content or len(content.strip()) < 30:
            return False
        
        # Conta palavras reais
        words = re.findall(r'\b\w{2,}\b', content)
        if len(words) < 8:
            return False
        
        # Verifica se não é só números ou códigos
        if re.match(r'^[\d\s\-\.,]+$', content):
            return False
        
        # Verifica diversidade lexical básica
        unique_words = set(word.lower() for word in words)
        if len(unique_words) < len(words) * 0.3:  # Muito repetitivo
            return False
        
        # Verifica se tem estrutura de frase
        if not re.search(r'\b\w+\s+\w+\b', content):  # Pelo menos duas palavras juntas
            return False
        
        return True

    def _optimize_chunk_boundaries(self, chunks: List[Dict]) -> List[Dict]:
        """Otimiza as fronteiras dos chunks para melhor coesão semântica"""
        if not chunks or len(chunks) < 2: # Adicionada checagem 'not chunks' para segurança
            return chunks if chunks is not None else [] # Retorna lista vazia se chunks for None
        
        optimized = []

        temp_chunks_status = [True] * len(chunks) # True se o chunk no índice i ainda precisa ser processado

        for i, current_chunk_original_ref in enumerate(chunks):
            if not temp_chunks_status[i]: # Se o chunk atual já foi mesclado com um anterior
                continue

            chunk = current_chunk_original_ref 


            content = chunk['conteudo'] # Se chunk fosse None, esta linha falharia.
                                        # Mas com temp_chunks_status, não deveríamos pegar um "consumido".

            # Usar .get() para 'palavras' para evitar KeyError se a chave não existir por algum motivo.
            word_count = chunk.get('palavras', len(content.split())) 


            if word_count < self.min_chunk_words and i < (len(chunks) - 1) and temp_chunks_status[i+1]:
                next_chunk_original_ref = chunks[i+1]

                combined_content_text = self._clean_chunk_text_optimized(content + ' ' + next_chunk_original_ref['conteudo'])
                combined_sentences = self.split_sentences(combined_content_text)
                combined_words = self._count_words_in_chunk(combined_sentences)

                if combined_words <= self.max_chunk_words * 1.2: # Limite de mesclagem
                    merged_chunk = chunk.copy() 
                    merged_chunk['conteudo'] = combined_content_text
                    merged_chunk['palavras'] = combined_words
                    merged_chunk['sentencas'] = len(combined_sentences) 
                    merged_chunk['tipo_quebra'] = 'mesclado'
                    # Calcula a confiança da mesclagem
                    conf_chunk = chunk.get('confianca', 0.8)
                    conf_next_chunk = next_chunk_original_ref.get('confianca', 0.8)
                    merged_chunk['confianca'] = float(min(conf_chunk, conf_next_chunk))
                    
                    optimized.append(merged_chunk)
                    temp_chunks_status[i] = False # Marca o chunk atual como processado (mesclado)
                    temp_chunks_status[i+1] = False # Marca o próximo chunk como processado (mesclado)
                    continue # Pula a adição individual do chunk atual
            
            # Se não mesclou (ou era o último, ou o próximo já foi processado, ou a mescla excedeu o tamanho)
            optimized.append(chunk)
            temp_chunks_status[i] = False # Marca como processado (adicionado individualmente)
            
        return optimized # 'optimized' não deve conter 'None' com esta lógica

    def _calculate_chunk_quality_score(self, chunk: Dict) -> float:
        """Calcula score de qualidade do chunk"""
        content = chunk.get('conteudo', '')
        if not content:
            return 0.0
        
        score = 1.0
        words = content.split()
        word_count = len(words)
        
        # Penaliza chunks muito pequenos ou muito grandes
        if word_count < self.min_chunk_words:
            score *= 0.5
        elif word_count > self.max_chunk_words:
            score *= 0.8
        
        # Bonifica chunks no tamanho ideal
        if self.min_chunk_words <= word_count <= self.ideal_chunk_words:
            score *= 1.1
        
        # Verifica terminação adequada
        if content.rstrip().endswith(('.', '!', '?')):
            score *= 1.05
        elif content.rstrip().endswith((':', ';')):
            score *= 1.02
        else:
            score *= 0.9
        
        # Verifica diversidade lexical
        unique_words = set(word.lower() for word in words if len(word) > 2)
        if words:
            diversity = len(unique_words) / len(words)
            score *= (0.7 + diversity * 0.6)  # Score entre 0.7 e 1.3
        
        # Usa confiança existing se disponível
        existing_confidence = chunk.get('confianca', 0.8)
        score = (score + existing_confidence) / 2
        
        return float(min(score, 1.0))


    def _try_split_single_oversized_chunk(self, oversized_chunk: Dict) -> List[Dict]:
        """
        Tenta dividir um chunk semanticamente se ele for muito longo.
        Retorna uma lista de chunks (1 se não dividido, 2+ se dividido).
        """
        content = oversized_chunk['conteudo']
        original_page = oversized_chunk.get('pagina')
        # Use as contagens originais como fallback se a divisão falhar
        original_palavras = oversized_chunk['palavras']
        original_sentencas = oversized_chunk['sentencas']

        internal_sentences = self.split_sentences(content)

   
        if len(internal_sentences) < self.min_sentences_per_chunk * 2:
            return [oversized_chunk]

     
        similarities = self._calculate_semantic_similarity_batch(internal_sentences)
        if not similarities: 
            return [oversized_chunk]

        best_split_point_idx = -1 
        lowest_similarity_at_split = 1.0

        for i in range(len(similarities)):
            num_sentences_chunk1 = i + 1
            num_sentences_chunk2 = len(internal_sentences) - num_sentences_chunk1

            if num_sentences_chunk1 < self.min_sentences_per_chunk or \
               num_sentences_chunk2 < self.min_sentences_per_chunk:
                continue

            sentences1 = internal_sentences[:num_sentences_chunk1]
            sentences2 = internal_sentences[num_sentences_chunk1:]

            words1 = self._count_words_in_chunk(sentences1)
            words2 = self._count_words_in_chunk(sentences2)

            if not (self.min_chunk_words <= words1 <= self.max_chunk_words and
                    words2 >= self.min_chunk_words):
                continue
            current_similarity = similarities[i]
            
            split_similarity_threshold = self.threshold_base - 0.15 

            if current_similarity < split_similarity_threshold:
                if current_similarity < lowest_similarity_at_split:
                    lowest_similarity_at_split = current_similarity
                    best_split_point_idx = i
        
        if best_split_point_idx != -1:
            split_after_sentence_index = best_split_point_idx
            chunk1_sentences = internal_sentences[:split_after_sentence_index + 1]
            chunk2_sentences = internal_sentences[split_after_sentence_index + 1:]

            chunk1_text = self._clean_chunk_text_optimized(' '.join(chunk1_sentences))
            chunk2_text = self._clean_chunk_text_optimized(' '.join(chunk2_sentences))
            
            # Revalida os novos chunks
            if not self._is_valid_semantic_chunk_words(chunk1_sentences) or \
               not self._is_valid_semantic_chunk_words(chunk2_sentences):
                return [oversized_chunk] # Divisão resultou em chunks inválidos

            palavras_c1 = self._count_words_in_chunk(chunk1_sentences)
            palavras_c2 = self._count_words_in_chunk(chunk2_sentences)

            new_chunks = [
                {
                    'pagina': original_page,
                    'conteudo': chunk1_text,
                    'tipo_quebra': 'split_long', # Indica que este chunk resultou de uma divisão
                    'confianca': float(max(0.6, min(0.95, 1.0 - lowest_similarity_at_split))), # Confiança baseada na dissimilaridade
                    'palavras': palavras_c1,
                    'sentencas': len(chunk1_sentences)
                },
                {
                    'pagina': original_page,
                    'conteudo': chunk2_text,
                    'tipo_quebra': 'split_long_cont', # Continuação do chunk dividido
                    'confianca': float(max(0.6, min(0.95, 1.0 - lowest_similarity_at_split))),
                    'palavras': palavras_c2,
                    'sentencas': len(chunk2_sentences)
                }
            ]
            print(f"✂️ Chunk longo (P:{original_page}, {original_palavras}p/{original_sentencas}s) dividido em 2. "
                  f"Novos: ({new_chunks[0]['palavras']}p/{new_chunks[0]['sentencas']}s) e ({new_chunks[1]['palavras']}p/{new_chunks[1]['sentencas']}s). "
                  f"Similaridade no split: {lowest_similarity_at_split:.2f}")
            return new_chunks
        
        return [oversized_chunk] # Não foi possível dividir ou encontrar um bom ponto de divisão

    def _split_all_oversized_chunks(self, chunks: List[Dict]) -> List[Dict]:
        """
        Itera sobre os chunks e tenta dividir aqueles que excedem max_chunk_words.
        Repete o processo se houver divisões, até um limite de iterações.
        """
        trigger_word_count = self.max_chunk_words 
        # Poderia usar um valor um pouco maior que self.max_chunk_words para dar uma margem, 
        # ex: self.max_chunk_words * 1.1, mas para atender ao pedido de quebrar os "muito longos",
        # usar self.max_chunk_words diretamente é mais alinhado.

        processed_chunks_in_iteration = True 
        current_chunks = list(chunks) 

        max_iterations = 5 # Para evitar loops infinitos em casos complexos
        num_iterations = 0

        while processed_chunks_in_iteration and num_iterations < max_iterations:
            processed_chunks_in_iteration = False
            num_iterations += 1
            
            next_iteration_chunks = []
            for chunk_item in current_chunks:
                # Verifica se o chunk atual precisa ser avaliado para divisão
                if chunk_item['palavras'] > trigger_word_count:
                    split_attempt_results = self._try_split_single_oversized_chunk(chunk_item)
                    next_iteration_chunks.extend(split_attempt_results)
                    if len(split_attempt_results) > 1: # Se o chunk foi efetivamente dividido
                        processed_chunks_in_iteration = True 
                else:
                    next_iteration_chunks.append(chunk_item) # Mantém o chunk como está
            current_chunks = next_iteration_chunks
            
            if processed_chunks_in_iteration:
                 print(f"🔄 Repassando divisão de chunks (Iteração {num_iterations}). Total atual: {len(current_chunks)} chunks.")

        return current_chunks



    def is_valid_content(self, content: str) -> bool:
        """Validação mais rigorosa com múltiplos critérios"""
        if not content or len(content) < self.min_words:
            return False
            
        # Verifica padrões inválidos
        invalid_patterns = [
            r'^\d+$',  # Apenas números
            r'^[^a-zA-Z0-9]{5,}$',  # Símbolos
            r'^\s*[A-Z]{2,}\s*$'  # Siglas
        ]
        
        if any(re.match(p, content) for p in invalid_patterns):
            return False
            
        # Verificação gramatical com fallback
        return self.tem_sentido(content) if len(content) < 120 else True

    def chunk_faq(self, lines: List[str]) -> List[Dict]:
        """Processa FAQ format"""
        chunks = []
        chunk_id = 1
        
        for i in range(len(lines) - 1):
            question_match = self.faq_q_re.match(lines[i])  # Nome correto
            answer_match = self.faq_a_re.match(lines[i + 1])  # Nome correto
            
            if question_match and answer_match:
                question = question_match.group(1).strip()
                answer = answer_match.group(1).strip()
                
                if question and answer:  # Valida se tem conteúdo
                    chunks.append({
                        'chunk_id': chunk_id,
                        'pergunta': question,
                        'resposta': answer
                    })
                    chunk_id += 1
        
        return chunks

    def parse_interview_blocks(self, lines: List[str]) -> List[tuple]:
        """Parse blocos de entrevista/diálogo"""
        blocks = []
        current_block = None
        current_header = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Detecta cabeçalho
            if self.header_interview_re.match(line):
                current_header = line
                continue
            
            # Detecta speaker
            speaker_match = self.speaker_re.match(line)
            if speaker_match:
                name, text = speaker_match.groups()
                name, text = name.strip(), text.strip()
                
                if current_block and current_block[0] == name:
                    current_block[1].append(text)
                else:
                    if current_block:
                        blocks.append((current_header, current_block))
                    current_block = [name, [text]]
        
        if current_block:
            blocks.append((current_header, current_block))
        
        return blocks

    def chunk_interview(self, blocks: List[tuple], has_header: bool = False) -> List[Dict]:
        """Processa formato de entrevista"""
        chunks = []
        chunk_id = 1
        
        # Processa pares de blocos (pergunta-resposta)
        for i in range(0, len(blocks) - 1, 2):
            header_a, block_a = blocks[i]
            header_b, block_b = blocks[i + 1]
            
            pergunta = f"{block_a[0]}: {' '.join(block_a[1])}"
            resposta = f"{block_b[0]}: {' '.join(block_b[1])}"
            
            chunk = {
                'chunk_id': chunk_id,
                'pergunta': pergunta,
                'resposta': resposta
            }
            
            if has_header and header_a:
                chunk['cabecalho'] = header_a
            
            chunks.append(chunk)
            chunk_id += 1
        
        return chunks

    def clear_cache(self):
        """Limpa caches para liberar memória"""
        self._validation_cache.clear()
        self._embedding_cache.clear()

    def chunk_semantic_v3_enhanced(self, path: str, progress_queue: Queue = None) -> List[Dict]:
        """Wrapper para compatibilidade com o sistema de progresso"""
        return self.chunk_structured(path)

    def _calculate_semantic_similarity_batch(self, sentences: List[str]) -> List[float]:
        """
        Calcula a similaridade semântica entre sentenças consecutivas em lote.
        """
        if len(sentences) < 2:
            return []

        # Garante que a biblioteca necessária está disponível
        if not ADVANCED_NLP:
            print("AVISO: Análise de similaridade requer bibliotecas avançadas.")
            # Retorna um valor padrão neutro para evitar quebra total
            return [0.8] * (len(sentences) - 1)

        try:
            # Usa o modelo principal para gerar os embeddings
            embeddings = self.model.encode(sentences, batch_size=self.batch_size, show_progress_bar=False)
            
            # Calcula a similaridade de cosseno entre pares de sentenças adjacentes
            # Compara o embedding[i] com o embedding[i+1]
            sims = cosine_similarity(
                embeddings[:-1],
                embeddings[1:]
            )
            
            # A similaridade entre a sentença 'i' e 'i+1' está na diagonal da matriz
            return [sims[i][i] for i in range(len(sims))]

        except Exception as e:
            print(f"⚠️ Erro ao calcular similaridade semântica: {e}")
            return [0.0] * (len(sentences) - 1)


class Application:
    def __init__(self, master):
        self.master = master
        self.processor = DocumentProcessor()
        self.json_data = None
        self.setup_ui()
        
    def setup_processor_with_api_key(self, api_key: str = None):
        """Configura o processador com chave da API"""
        self.processor = DocumentProcessor(gemini_api_key=api_key)

    def setup_ui(self):
        style = ttk.Style()
        style.theme_use('clam')
        
        # Cores do tema dark (mantidas as mesmas)
        dark_bg = '#23272e'
        dark_panel = '#2c313a'
        accent = '#5e81ac'
        accent_hover = '#81a1c1'
        text_color = '#e5e9f0'
        label_color = '#bfc9db'
        entry_bg = '#181a20'
        border_color = '#444857'
        card_bg = '#252a33'
        card_border = '#3b4252'
        shadow = '#181a20'

        # Configuração de estilos (mantida igual)
        style.configure('TFrame', background=dark_bg)
        style.configure('TLabel', background=dark_bg, foreground=label_color, font=('Segoe UI', 11))
        style.configure('Title.TLabel', background=dark_bg, foreground=accent, font=('Segoe UI', 20, 'bold'))
        style.configure('Subtitle.TLabel', background=dark_bg, foreground=label_color, font=('Segoe UI', 12, 'italic'))
        style.configure('TButton', font=('Segoe UI', 12, 'bold'), padding=10, background=accent, foreground=text_color, borderwidth=0)
        style.map('TButton', background=[('active', accent_hover)], foreground=[('active', text_color)])
        style.configure('TCheckbutton', background=dark_bg, foreground=label_color, font=('Segoe UI', 10))
        style.configure('TEntry', fieldbackground=entry_bg, foreground=text_color, bordercolor=border_color, font=('Segoe UI', 12))
        style.configure('TCombobox', fieldbackground=entry_bg, foreground=text_color, background=entry_bg, bordercolor=border_color, font=('Segoe UI', 12))
        style.map('TCombobox', fieldbackground=[('readonly', entry_bg)], foreground=[('readonly', text_color)])

        self.master.title("Processador de Documentos - Versão Melhorada")
        self.master.geometry("950x750")
        self.master.configure(bg=dark_bg)

        # Interface (mantida similar, mas com melhorias)
        title = ttk.Label(self.master, text="Processador de Documentos", style='Title.TLabel')
        title.pack(pady=(18, 2))
        subtitle = ttk.Label(self.master, text="Chunking semântico inteligente com validação de conteúdo", style='Subtitle.TLabel')
        subtitle.pack(pady=(0, 12))

        # Frame principal
        main_frame = ttk.Frame(self.master, padding=24, style='TFrame')
        main_frame.pack(fill=X, padx=40, pady=(0, 10))

        # Linha do arquivo
        file_frame = ttk.Frame(main_frame, style='TFrame')
        file_frame.pack(fill=X, pady=10)
        ttk.Label(file_frame, text="Arquivo:").pack(side=LEFT, padx=(0, 8))
        self.file_entry = ttk.Entry(file_frame, width=48)
        self.file_entry.pack(side=LEFT, padx=(0, 8))
        btn_browse = ttk.Button(file_frame, text="🔍 Procurar", command=self.browse_file)
        btn_browse.pack(side=LEFT)
        download_btn = ttk.Button(file_frame, text="⬇️ Baixar JSON", command=self.save_json)
        download_btn.pack(side=LEFT, padx=(12, 0))

        # ==== Campo para API Key do Gemini ====
        api_frame = ttk.Frame(main_frame, style='TFrame')
        api_frame.pack(fill=X, pady=10)
        ttk.Label(api_frame, text="API Key Gemini:").pack(side=LEFT, padx=(0, 8))
        self.api_key_entry = ttk.Entry(api_frame, width=40, show="*")
        self.api_key_entry.pack(side=LEFT, padx=(0, 8))
        btn_config_api = ttk.Button(api_frame, text="🔑 Configurar", command=self.configure_gemini_api)
        btn_config_api.pack(side=LEFT)
        self.api_status_label = ttk.Label(api_frame, text="❌ Gemini não configurado", 
                                         font=('Segoe UI', 9, 'italic'))
        self.api_status_label.pack(side=LEFT, padx=(12, 0))

        # Configurações avançadas
        config_frame = ttk.Frame(main_frame, style='TFrame')
        config_frame.pack(fill=X, pady=10)
        
        ttk.Label(config_frame, text="Estilo:").pack(side=LEFT, padx=(0, 8))
        self.doc_type = ttk.Combobox(config_frame, values=["FAQ", "Pergunta-Resposta/Entrevista", "Texto Puro"], 
                                   state="readonly", width=24)
        self.doc_type.pack(side=LEFT, padx=(0, 8))
        self.doc_type.bind("<<ComboboxSelected>>", self.toggle_header_check)
        
        self.header_var = IntVar()
        self.header_check = ttk.Checkbutton(config_frame, text="Contém cabeçalhos", variable=self.header_var)
        
        # Controles de threshold
        ttk.Label(config_frame, text="Similaridade:").pack(side=LEFT, padx=(15, 5))
        self.threshold_var = DoubleVar(value=0.75)
        threshold_scale = ttk.Scale(config_frame, from_=0.5, to=0.9, variable=self.threshold_var, 
                                  orient=HORIZONTAL, length=100)
        threshold_scale.pack(side=LEFT, padx=(0, 5))
        self.threshold_label = ttk.Label(config_frame, text="0.75")
        self.threshold_label.pack(side=LEFT)
        threshold_scale.configure(command=self.update_threshold_label)

        # Status e progresso
        self.status_var = StringVar(value="Pronto para processar")
        status_label = ttk.Label(main_frame, textvariable=self.status_var, 
                               font=('Segoe UI', 10, 'italic'))
        status_label.pack(pady=(10, 0))

        # Card de exemplo (mantido)
        card_frame = Frame(main_frame, bg=card_bg, highlightbackground=card_border, highlightthickness=1, bd=0)
        card_frame.pack(fill=X, pady=(10, 16), padx=2)
        self.example_label = Label(card_frame, text="", font=("Segoe UI", 11, "italic"), 
                                 fg=accent, bg=card_bg, justify=LEFT, anchor=W, padx=12, pady=8)
        self.example_label.pack(fill=X, anchor=W)
        self.update_example_label()

        # Botão de processar
        self.process_btn = ttk.Button(main_frame, text="⚡ Processar Arquivo", command=self.process_file)
        self.process_btn.pack(pady=(0, 10), ipadx=10)

        # Frame de visualização
        view_frame = Frame(self.master, bg=shadow)
        view_frame.pack(fill=BOTH, expand=1, padx=40, pady=(0, 10))
        result_card = Frame(view_frame, bg=card_bg, highlightbackground=card_border, highlightthickness=2, bd=0)
        result_card.pack(fill=BOTH, expand=1, padx=0, pady=0)
        
        # Header do resultado
        result_header = ttk.Frame(result_card, style='TFrame')
        result_header.configure(style='Card.TFrame')
        result_header.pack(fill=X, pady=(8, 0), padx=12)
        
        ttk.Label(result_header, text="Resultado:", font=("Segoe UI", 12, "bold"), 
                foreground=accent, background=card_bg).pack(side=LEFT)
        
        self.chunk_count_label = ttk.Label(result_header, text="", font=("Segoe UI", 10), 
                                         foreground=label_color, background=card_bg)
        self.chunk_count_label.pack(side=RIGHT)
        
        # Área de texto
        text_frame = Frame(result_card, bg=card_bg)
        text_frame.pack(fill=BOTH, expand=1, padx=12, pady=(5, 12))
        
        self.json_text = Text(text_frame, wrap=NONE, font=("Consolas", 11), 
                            bg=dark_panel, fg=text_color, insertbackground=accent, 
                            relief=GROOVE, borderwidth=2, height=18)
        self.json_text.pack(side=LEFT, fill=BOTH, expand=1)
        
        scroll_y = ttk.Scrollbar(text_frame, command=self.json_text.yview)
        scroll_y.pack(side=RIGHT, fill=Y)
        self.json_text.configure(yscrollcommand=scroll_y.set)

        # --- Barra de progresso ---
        self.progress_bar = ttk.Progressbar(main_frame, orient='horizontal', length=300, mode='determinate')
        self.progress_bar.pack(pady=(5, 10), fill=X, padx=50)
        self.progress_bar.pack_forget()  # Esconde inicialmente

    def update_threshold_label(self, value):
        """Atualiza label do threshold"""
        self.threshold_label.config(text=f"{float(value):.2f}")
        self.processor.threshold = float(value)

    def toggle_header_check(self, event=None):
        """Toggle checkbox de cabeçalho"""
        if self.doc_type.get() == "Pergunta-Resposta/Entrevista":
            self.header_check.pack(side=LEFT, padx=(8, 0))
        else:
            self.header_check.pack_forget()
        self.update_example_label()

    def update_example_label(self):
        """Atualiza exemplo baseado no tipo selecionado"""
        estilo = self.doc_type.get()
        if estilo == "FAQ":
            exemplo = "💡 Exemplo:\n1. Qual é a capital do Brasil?\nR: Brasília é a capital."
        elif estilo == "Pergunta-Resposta/Entrevista":
            exemplo = "💡 Exemplo:\nEntrevistador: Como você começou?\nEntrevistado: Comecei há 5 anos..."
        else:
            exemplo = "💡 Exemplo:\nTexto livre será dividido em chunks semânticos"
            if hasattr(self, 'processor') and self.processor and self.processor.use_gemini:
                exemplo += " usando IA Gemini para máxima precisão! 🤖"
            else:
                exemplo += " usando embeddings tradicionais."
        self.example_label.config(text=exemplo)

    def browse_file(self):
        """Seleciona arquivo"""
        filepath = filedialog.askopenfilename(
            title="Selecionar documento",
            filetypes=[
                ("Todos os suportados", "*.txt *.pdf *.docx"),
                ("Arquivos de texto", "*.txt"),
                ("PDFs", "*.pdf"),
                ("Word", "*.docx")
            ]
        )
        if filepath:
            self.file_entry.delete(0, END)
            self.file_entry.insert(0, filepath)
            self.status_var.set(f"Arquivo selecionado: {os.path.basename(filepath)}")

    def process_file(self):
        """Inicia o processamento em thread separada"""
        path = self.file_entry.get().strip()
        if not path or not os.path.isfile(path):
            messagebox.showerror("Erro", "Selecione um caminho de arquivo válido!")
            return
        if not self.processor:
            messagebox.showerror("Erro", "O processador de documentos não foi inicializado.")
            return
        # Prepara a UI
        self.process_btn.config(state="disabled")
        self.status_var.set("Iniciando processamento...")
        self.json_text.delete(1.0, END)
        self.chunk_count_label.config(text="")
        self.progress_bar.pack()
        self.progress_bar['value'] = 0
        self.master.update_idletasks()
        # Cria queue e thread
        self.result_queue = Queue()
        doc_type = self.doc_type.get()
        has_header = self.header_var.get() == 1 if doc_type == "Pergunta-Resposta/Entrevista" else False
        threading.Thread(
            target=self._background_process_file,
            args=(path, doc_type, has_header),
            daemon=True
        ).start()
        # Inicia verificação da queue
        self.master.after(100, self._check_processing_queue)

    def _background_process_file(self, path, doc_type, has_header):
        """Processamento em background"""
        try:
            if doc_type == "FAQ":
                # Passe is_faq_mode=True para o load_lines
                lines = self.processor.load_lines(path, is_faq_mode=True) 
                # Opcional: Adicione um print para depuração
                # print(f"DEBUG: Linhas para FAQ (total {len(lines)}):")
                # for i, l in enumerate(lines):
                #     print(f"  {i}: {l}")
                chunks = self.processor.chunk_faq(lines)
                self.result_queue.put({"type": "result", "success": True, "data": chunks})
            elif doc_type == "Pergunta-Resposta/Entrevista":
                # Para outros tipos, chame load_lines sem is_faq_mode (ou com False)
                lines = self.processor.load_lines(path) 
                blocks = self.processor.parse_interview_blocks(lines)
                chunks = self.processor.chunk_interview(blocks, has_header)
                self.result_queue.put({"type": "result", "success": True, "data": chunks})
            else:  # Texto Puro
                # chunk_structured não usa load_lines diretamente da mesma forma,
                # ele usa load_document. Não precisa de is_faq_mode aqui.
                chunks = self.processor.chunk_structured(path)
                self.result_queue.put({"type": "result", "success": True, "data": chunks})
        except Exception as e:
            # print(f"DEBUG: Erro em _background_process_file: {e}") # Para depuração
            self.result_queue.put({"type": "result", "success": False, "error": str(e)})

    def _check_processing_queue(self):
        """Verifica atualizações na queue"""
        try:
            while not self.result_queue.empty():
                msg = self.result_queue.get_nowait()
                if msg['type'] == 'progress':
                    self.progress_bar['value'] = msg['value']
                    self.status_var.set(msg['step'])
                    self.master.update_idletasks()
                elif msg['type'] == 'result':
                    self.progress_bar.pack_forget()
                    self.process_btn.config(state="normal")
                    if msg['success']:
                        chunks = msg['data']
                        self.json_data = json.dumps(chunks, ensure_ascii=False, indent=2)
                        self.json_text.insert(END, self.json_data)
                        self.chunk_count_label.config(text=f"{len(chunks)} chunks")
                        messagebox.showinfo("Sucesso", f"Processamento concluído!\n{len(chunks)} chunks gerados.")
                    else:
                        messagebox.showerror("Erro", f"Falha no processamento:\n{msg['error']}")
                    return
            self.master.after(100, self._check_processing_queue)
        except Exception as e:
            logging.error(f"Erro na queue: {str(e)}")
            self.progress_bar.pack_forget()
            self.process_btn.config(state="normal")

    def save_json(self):
        """Salva o JSON processado"""
        if not self.json_data:
            messagebox.showwarning("Aviso", "Nenhum dado processado para salvar!")
            return

        entrada = self.file_entry.get()
        if entrada:
            base = os.path.splitext(os.path.basename(entrada))[0]
            sugestao = f"{base}_chunks.json"
        else:
            sugestao = "chunks.json"

        filepath = filedialog.asksaveasfilename(
            title="Salvar chunks como JSON",
            defaultextension=".json",
            initialfile=sugestao,
            filetypes=[("JSON files", "*.json"), ("Todos os arquivos", "*.")]
        )
        
        if filepath:
            try:
                with open(filepath, 'w', encoding='utf-8') as f:
                    try:
                        json.loads(self.json_data)  # Valida o JSON
                        f.write(self.json_data)
                    except Exception as e:
                        messagebox.showerror("Erro", f"JSON inválido: {str(e)}")
                        return
                messagebox.showinfo("Sucesso", f"Arquivo salvo em:\n{filepath}")
                self.status_var.set(f"Salvo: {os.path.basename(filepath)}")
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao salvar arquivo:\n{str(e)}")

    def configure_gemini_api(self):
        """Configura a API do Gemini de forma não-bloqueante"""
        api_key = self.api_key_entry.get().strip()
        if not api_key:
            messagebox.showwarning("Aviso", "Digite uma chave API válida!")
            return
        
        # Desabilita controles durante teste
        self.api_key_entry.config(state="disabled")
        self.api_status_label.config(text="🔄 Testando...", foreground='#d08770')
        self.master.update_idletasks()
        
        def test_api():
            try:
                genai.configure(api_key=api_key)
                # Teste mais rápido e simples
                test_model = genai.GenerativeModel('gemini-1.5-flash')
                test_response = test_model.generate_content("Responda apenas: OK")
                
                # Configura processador
                self.setup_processor_with_api_key(api_key)
                
                # Atualiza UI na thread principal
                self.master.after(0, lambda: self._api_success())
                
            except Exception as e:
                error_msg = str(e)
                self.master.after(0, lambda: self._api_error(error_msg))
        
        # Executa teste em thread separada
        threading.Thread(target=test_api, daemon=True).start()

    def _api_success(self):
        """Callback para sucesso na API"""
        self.api_status_label.config(text="✅ Gemini ativo", foreground='#5e81ac')
        self.api_key_entry.config(state="normal")
        messagebox.showinfo("Sucesso", "🤖 Gemini configurado!\nChunking semântico com IA ativado.")
        self.update_example_label()  # Atualiza exemplo

    def _api_error(self, error_msg):
        """Callback para erro na API"""
        self.api_status_label.config(text="❌ Erro na chave", foreground='#bf616a')
        self.api_key_entry.config(state="normal")
        messagebox.showerror("Erro", f"Falha ao configurar Gemini:\n{error_msg}\n\nVerifique a chave e créditos.")


def main():
    """Função principal"""
    # Configura logging
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
    
    # Suprime logs verbosos do pdfminer
    logging.getLogger("pdfminer").setLevel(logging.ERROR)
    
    # Cria e executa aplicação
    root = Tk()
    app = Application(root)
    
    try:
        root.mainloop()
    except KeyboardInterrupt:
        print("Aplicação interrompida pelo usuário")
    except Exception as e:
        logging.error(f"Erro fatal na aplicação: {e}")
        messagebox.showerror("Erro Fatal", f"Erro inesperado:\n{str(e)}")


if __name__ == '__main__':
    main()
